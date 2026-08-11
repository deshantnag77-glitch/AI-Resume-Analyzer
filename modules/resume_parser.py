"""
resume_parser.py
-----------------
Handles extraction of raw text from uploaded resume files.
Supports PDF, DOCX and TXT formats.
"""

import io
from pypdf import PdfReader
import docx


class ResumeParseError(Exception):
    """Raised when a resume file cannot be parsed or contains no usable text."""
    pass


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract raw text from a PDF file given as bytes."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        if len(reader.pages) == 0:
            raise ResumeParseError("The PDF file has no pages.")

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        text = "\n".join(text_parts).strip()

        if not text:
            raise ResumeParseError(
                "No readable text found in the PDF. "
                "This may be a scanned/image-based resume."
            )
        return text

    except ResumeParseError:
        raise
    except Exception as e:
        raise ResumeParseError(f"Could not read the PDF file: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract raw text from a DOCX file given as bytes."""
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        # Also pull text from tables, since some resumes use table layouts
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        text = "\n".join(paragraphs).strip()

        if not text:
            raise ResumeParseError("No readable text found in the DOCX file.")
        return text

    except ResumeParseError:
        raise
    except Exception as e:
        raise ResumeParseError(f"Could not read the DOCX file: {e}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from a plain TXT file given as bytes."""
    try:
        text = file_bytes.decode("utf-8", errors="ignore").strip()
        if not text:
            raise ResumeParseError("The TXT file is empty.")
        return text
    except ResumeParseError:
        raise
    except Exception as e:
        raise ResumeParseError(f"Could not read the TXT file: {e}")


def parse_resume(uploaded_file) -> str:
    """
    Main entry point. Takes a Streamlit UploadedFile object,
    detects its type from the filename, and returns extracted text.
    """
    if uploaded_file is None:
        raise ResumeParseError("No file was uploaded.")

    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()

    if not file_bytes:
        raise ResumeParseError("The uploaded file is empty.")

    if filename.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    elif filename.endswith(".txt"):
        text = extract_text_from_txt(file_bytes)
    else:
        raise ResumeParseError(
            "Unsupported file format. Please upload a PDF, DOCX or TXT file."
        )

    if len(text.split()) < 20:
        raise ResumeParseError(
            "The resume text seems too short to analyze. "
            "Please check the file and try again."
        )

    return text
